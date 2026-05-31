from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("IIITL_Mess_Portal_Project_Report.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="8C8C8C", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:{}".format(m)))
        if node is None:
            node = OxmlElement("w:{}".format(m))
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_3 = OxmlElement("w:fldChar")
    fld_char_3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    run._r.append(text)
    run._r.append(fld_char_3)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 18, "000000"),
        ("Heading 1", 16, "000000"),
        ("Heading 2", 14, "000000"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def add_footer_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")


def add_centered_paragraph(doc, text="", size=12, bold=False, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_fill_line(doc, label, width_hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("[FILL: {}{}]".format(label, " - " + width_hint if width_hint else ""))
    run.bold = True
    run.font.color.rgb = RGBColor(120, 0, 0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "E8EEF5")
        set_cell_border(hdr[idx], "6E7F91", "8")
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in hdr[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx], "B7B7B7", "6")
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx == 0 and len(headers) > 2:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def add_placeholder_box(doc, label, caption=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.1)
    set_cell_border(cell, "7F7F7F", "12")
    set_cell_shading(cell, "F2F2F2")
    set_cell_margins(cell, top=360, bottom=360, start=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[IMAGE PLACEHOLDER: {}]".format(label))
    r.bold = True
    r.font.color.rgb = RGBColor(90, 90, 90)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(caption).italic = True
    doc.add_paragraph()


def add_report_image(doc, image_path, caption):
    path = Path(image_path)
    if not path.exists():
        add_placeholder_box(doc, "Missing image: {}".format(image_path), caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).italic = True
    doc.add_paragraph()


def chapter_title(doc, number, title):
    doc.add_page_break()
    add_centered_paragraph(doc, "Chapter {}".format(number), size=16, bold=True, spacing_after=8)
    add_centered_paragraph(doc, title, size=18, bold=True, spacing_after=18)


def add_blank_page(doc):
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    add_footer_page_number(doc.sections[0])

    # Title page
    for _ in range(5):
        doc.add_paragraph()
    add_centered_paragraph(doc, "IIITL MESS PORTAL", size=20, bold=True, spacing_after=16)
    add_centered_paragraph(doc, "A project report submitted in partial fulfillment of the requirements for the", size=12)
    add_centered_paragraph(doc, "award of the degree of", size=12)
    add_centered_paragraph(doc, "B.Tech. in Computer Science", size=14, bold=True, spacing_after=18)
    add_centered_paragraph(doc, "by", size=12)
    add_centered_paragraph(doc, "[FILL: Candidate Name]", size=13, bold=True)
    add_centered_paragraph(doc, "([FILL: Roll Number])", size=12, spacing_after=18)
    add_centered_paragraph(doc, "under the guidance of", size=12)
    add_centered_paragraph(doc, "[FILL: Supervisor Name]", size=13, bold=True, spacing_after=22)
    add_centered_paragraph(doc, "Indian Institute of Information Technology, Lucknow", size=13, bold=True)
    add_centered_paragraph(doc, "[FILL: Month Year, e.g., May 2026]", size=12, spacing_after=28)
    add_centered_paragraph(doc, "(c) Indian Institute of Information Technology, Lucknow [FILL: Year].", size=10)
    add_blank_page(doc)

    # Declaration
    add_centered_paragraph(doc, "Declaration of Authorship", size=16, bold=True, spacing_after=18)
    p = doc.add_paragraph()
    p.add_run("I/we, ").bold = False
    p.add_run("[FILL: Your Name(s)]").bold = True
    p.add_run(", declare that the work presented in ")
    p.add_run('"IIITL Mess Portal"').italic = True
    p.add_run(" is my/our own. I/we confirm that:")
    add_bullets(doc, [
        "This work was completed entirely while in candidature for the B.Tech. degree at Indian Institute of Information Technology, Lucknow.",
        "Where I/we have consulted the published work of others, it is always cited.",
        "Wherever I/we have cited the work of others, the source is always indicated. Except for the cited material, this work is solely my/our work.",
        "I/we have acknowledged all major sources of information used during analysis, design, implementation, and testing.",
        "The screenshots and deployment details included in the final submission will correspond to the implemented system."
    ])
    doc.add_paragraph()
    doc.add_paragraph("Signed:")
    add_table(doc, ["Author 1", "Author 2", "Author 3", "Author 4"], [["[FILL]", "[FILL]", "[FILL]", "[FILL]"]], [1.5, 1.5, 1.5, 1.5])
    add_fill_line(doc, "Date")
    add_blank_page(doc)

    # Certificate
    add_centered_paragraph(doc, "CERTIFICATE", size=16, bold=True, spacing_after=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("This is to certify that the work entitled ")
    p.add_run('"IIITL Mess Portal"').bold = True
    p.add_run(" submitted by ")
    p.add_run("[FILL: Your Name]").bold = True
    p.add_run(" who got his/her name registered on ")
    p.add_run("[FILL: Registration Month Year]").bold = True
    p.add_run(" for the award of B.Tech. degree at Indian Institute of Information Technology, Lucknow is based upon his/her own work under the supervision of ")
    p.add_run("[FILL: Supervisor Name, Department, Institute]").bold = True
    p.add_run(" and that neither this work nor any part of it has been submitted for any degree/diploma or any other academic award anywhere before.")
    for _ in range(4):
        doc.add_paragraph()
    add_fill_line(doc, "Name of Your Supervisor")
    add_fill_line(doc, "Department")
    add_fill_line(doc, "Institute/University")
    add_fill_line(doc, "Pin Code, INDIA")
    add_blank_page(doc)

    # Acknowledgements
    add_centered_paragraph(doc, "Acknowledgements", size=16, bold=True, spacing_after=18)
    doc.add_paragraph(
        "I would like to express my sincere gratitude to [FILL: Supervisor Name] for guidance, feedback, and encouragement throughout the development of this project. "
        "I am also thankful to the faculty members of Indian Institute of Information Technology, Lucknow for providing the academic environment and technical foundation required to complete this work."
    )
    doc.add_paragraph(
        "I also thank the mess administration and student users whose practical requirements shaped the design of this portal. Their needs helped define the core workflows of weekly meal pre-booking, QR-based meal verification, menu management, meal headcount estimation, feedback, and reminder notifications."
    )
    doc.add_paragraph(
        "Finally, I thank my family, friends, and peers for their continuous support during the implementation, testing, and documentation of this project."
    )
    doc.add_paragraph()
    add_fill_line(doc, "Lucknow")
    add_fill_line(doc, "Your name")
    add_fill_line(doc, "Month Year")
    add_blank_page(doc)

    # Abstract
    add_centered_paragraph(doc, "ABSTRACT", size=16, bold=True, spacing_after=18)
    doc.add_paragraph(
        "The IIITL Mess Portal is a web-based system made for a common problem in college mess management. In the present semester payment model, students usually pay mess fees for the whole semester even when they do not eat every meal. "
        "Because of holidays, travel, classes, illness, internships, or personal reasons, many meals are missed, but the student has already paid for them. This means some part of the student's money is wasted. "
        "The mess also faces a problem because it does not know the actual meal-wise headcount before cooking. Due to this, extra food may be prepared, food may get wasted, or sometimes the prepared quantity may not match the crowd."
    )
    doc.add_paragraph(
        "The proposed portal allows students to buy or pre-book meals for the upcoming week instead of depending only on a fixed semester payment. Students can view the menu, select meals, use QR codes for verification, check purchase history, and rate meals. "
        "The system also includes an AI-based meal selection feature using the Gemini API. A student can enter diet preferences, such as avoiding fried items, and the system suggests suitable meals from the weekly menu. "
        "For the mess side, the portal gives a clearer headcount before food preparation. This helps reduce wastage, makes resource planning easier, and helps control overcrowding because the mess can prepare food according to expected demand."
    )
    add_blank_page(doc)

    # Contents
    add_centered_paragraph(doc, "Contents", size=16, bold=True, spacing_after=14)
    toc_rows = [
        ["1", "Introduction", "[UPDATE PAGE NO.]"],
        ["2", "Literature Review", "[UPDATE PAGE NO.]"],
        ["3", "Methodology", "[UPDATE PAGE NO.]"],
        ["4", "Simulation and Results", "[UPDATE PAGE NO.]"],
        ["5", "Conclusion and Future Work", "[UPDATE PAGE NO.]"],
        ["Appendix", "Appendix Title Here", "[UPDATE PAGE NO.]"],
        ["", "Bibliography", "[UPDATE PAGE NO.]"],
    ]
    add_table(doc, ["Chapter", "Title", "Page"], toc_rows, [1.0, 4.2, 1.0])
    doc.add_paragraph("Note: Right-click and update this table manually if your institute requires exact page numbers after adding screenshots.")
    add_blank_page(doc)

    # Chapter 1
    chapter_title(doc, 1, "Introduction")
    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph(
        "A college mess serves many students every day, so planning the food quantity is very important. In many colleges, the mess fee is collected as a one-time semester payment. This is simple for fee collection, but it does not match the actual number of meals a student eats. "
        "A student may miss several meals during the semester, but the payment has already been made. On the other hand, the mess may still prepare food assuming a larger crowd."
    )
    doc.add_paragraph(
        "The IIITL Mess Portal is designed as a centralized web application for the Indian Institute of Information Technology, Lucknow mess workflow. "
        "It gives students a single place to view meal schedules, buy or pre-book next-week meals, auto-select meals based on diet preferences, access QR codes, and provide dish feedback. It gives administrators tools to update menu data, verify booked meals, analyze expected meal counts, and send automated reminders."
    )
    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "The existing semester-based mess payment approach creates problems for both students and the mess. Students lose money for meals they do not eat. The mess does not get a clear meal-wise headcount before cooking, so it becomes difficult to decide how much food to prepare. "
        "If too much food is prepared, wastage becomes a major issue. If less food is prepared, the mess can become overcrowded and some items may run out. The proposed system reduces these problems by allowing students to buy meals for the upcoming week and by giving the mess administration a clearer headcount before preparation."
    )
    doc.add_heading("1.3 Objectives", level=2)
    add_bullets(doc, [
        "To build a student-facing portal for viewing menus, timings, booked meal status, QR code, and purchase history.",
        "To implement secure Google OAuth login so that only authenticated users can access user-level services.",
        "To replace fixed semester-only payment dependency with weekly meal purchase or pre-booking using Razorpay order creation and payment verification.",
        "To provide AI-based meal selection using Gemini API according to student diet preferences such as avoiding fried items.",
        "To provide QR-code-based meal validation for mess service access.",
        "To provide administrative features for menu management, expected meal count analysis, dish rating review, and QR scanning.",
        "To help the mess reduce food wastage, control overcrowding, and manage resources based on advance meal demand.",
        "To send automated email reminders to users who have not bought next-week meals before the weekly deadline.",
        "To automatically roll over next-week meal bookings into current-week valid bookings at the beginning of the week."
    ])
    doc.add_heading("1.4 Scope of the Project", level=2)
    doc.add_paragraph(
        "The project focuses on the core operational requirements of a digital mess meal booking system: authentication, menu viewing, weekly meal purchase, AI-based meal selection, payment verification, QR-based meal validation, admin management, reminders, headcount analysis, and weekly rollover. "
        "The current implementation is suitable for a campus deployment where users log in through Google accounts and administrators manage mess data from a protected admin panel."
    )
    doc.add_heading("1.5 Project Overview", level=2)
    add_table(doc, ["Area", "Description"], [
        ["Project Name", "IIITL Mess Portal"],
        ["Frontend", "React-based single-page application"],
        ["Backend", "Node.js and Express REST API"],
        ["Database", "MongoDB using Mongoose models"],
        ["Authentication", "Passport Google OAuth 2.0"],
        ["Payment", "Razorpay order and signature verification"],
        ["AI Meal Selection", "Gemini API based selection using student diet preferences"],
        ["Notification", "Nodemailer SMTP email reminder service"],
        ["Automation", "Scheduled meal purchase reminder and weekly rollover services"],
    ], [1.8, 4.4])
    add_placeholder_box(doc, "Insert screenshot of homepage or dashboard", "Figure 1.1: Homepage or dashboard placeholder")

    # Chapter 2
    chapter_title(doc, 2, "Literature Review")
    doc.add_heading("2.1 Existing Semester Payment Model", level=2)
    doc.add_paragraph(
        "In a semester payment model, students pay mess fees for a long fixed period irrespective of their actual meal consumption. This model is administratively simple but not student-friendly when students are absent, travel home, skip meals, or are unavailable because of academic or personal commitments. "
        "It also does not provide the mess with a reliable meal-wise headcount before food preparation."
    )
    doc.add_heading("2.2 Digital Cafeteria and Mess Management Systems", level=2)
    doc.add_paragraph(
        "Digital cafeteria systems usually provide menu display, payment options, order history, and admin dashboards. A college mess is slightly different from a normal cafeteria because food is prepared in bulk and the same service repeats every day. "
        "So the important requirement is not only taking payment, but also knowing how many students are expected for each meal. A weekly meal booking model is useful because it gives students flexibility and also gives the mess useful planning data."
    )
    doc.add_heading("2.3 AI-Based Meal Recommendation", level=2)
    doc.add_paragraph(
        "Students may have different diet preferences. For example, some students may want to avoid fried food, some may prefer lighter meals, and some may want to skip certain dishes. Manually checking the whole weekly menu every time can be inconvenient. "
        "The project uses the Gemini API to help with this. The student enters diet preferences, and the system reads the weekly menu and selects suitable meals in JSON format. This makes meal selection faster and more personal."
    )
    doc.add_heading("2.4 QR-Based Verification", level=2)
    doc.add_paragraph(
        "QR-based verification is widely used because it is quick, inexpensive, and compatible with mobile devices. In the mess portal, each buyer has a secret code associated with their email. "
        "At serving time, the QR data can be checked against the active week and meal type. Once a booked meal is consumed, its corresponding value is updated to prevent reuse."
    )
    doc.add_heading("2.5 Online Payment and Notification Systems", level=2)
    doc.add_paragraph(
        "Online payment gateways reduce cash handling and provide structured payment confirmation through order identifiers and signatures. Notification systems such as email reminders support deadline-driven workflows by alerting users before a service cutoff. "
        "In this project, payment confirmation updates next-week meal booking data, while the reminder service sends an email when a user has not bought next-week meals by the reminder time."
    )
    doc.add_heading("2.6 Research Gap and Project Relevance", level=2)
    add_bullets(doc, [
        "A mess system needs flexible meal booking instead of a fixed semester-only payment dependency.",
        "Students should pay for selected meals rather than losing money for meals they do not consume.",
        "The mess needs aggregate meal counts before preparation begins so food wastage can be reduced.",
        "Advance headcount helps the mess manage ingredients, staff effort, and serving counters in a better way.",
        "Better preparation estimates can also help reduce overcrowding because popular items can be prepared in suitable quantity.",
        "AI-based selection can help students choose meals according to their diet preferences without manually checking every item.",
        "Students need reminders before losing eligibility for the next week.",
        "Meal verification should be fast enough for real serving-time use.",
        "The system must combine authentication, payment, QR validation, and administration in one flow."
    ])

    # Chapter 3
    chapter_title(doc, 3, "Methodology")
    doc.add_heading("3.1 Requirement Analysis", level=2)
    doc.add_paragraph(
        "The requirements were divided into student requirements, administrator requirements, and system automation requirements. Students need to log in, view menus, buy selected meals for the upcoming week, use AI suggestions based on diet preferences, check QR codes, and view purchase status. "
        "Administrators need to update timings and menu, scan QR codes, review expected meal counts, and view dish ratings. The system also needs scheduled jobs for reminder emails and weekly rollover of booked meals."
    )
    add_table(doc, ["Stakeholder", "Functional Requirements"], [
        ["Student", "Login, view menu, use AI meal selection, buy selected meals, view QR code, view purchase history, rate meals"],
        ["Administrator", "Set menu, set meal timings/costs, scan QR codes, view expected meal counts, view ratings"],
        ["System", "Verify payment, update booked meal state, send reminder email, roll over weekly bookings"],
    ], [1.6, 4.6])
    doc.add_heading("3.2 System Architecture", level=2)
    doc.add_paragraph(
        "The architecture follows a client-server model. The React frontend communicates with Express API routes. The backend uses Passport for authentication, Mongoose for database operations, Razorpay for payment processing, Gemini API for diet-preference-based meal selection, and Nodemailer for sending reminder emails. "
        "MongoDB stores users, buyer meal booking status, orders, menu records, meal timing and cost data, dish ratings, and scheduler logs."
    )
    add_report_image(doc, "report_assets/architecture_diagram.png", "Figure 3.1: System architecture of IIITL Mess Portal")
    doc.add_heading("3.3 Database Design", level=2)
    add_table(doc, ["Model", "Purpose", "Important Fields"], [
        ["User", "Stores authenticated user profile", "googleId, displayName, email"],
        ["Buyer", "Stores meal booking status and QR secret", "email, secret, bought, this, next"],
        ["Order", "Temporarily stores Razorpay order selection", "orderid, selected"],
        ["Menu", "Stores weekly menu", "day, breakfast, lunch, dinner"],
        ["Time", "Stores meal timing and cost", "meal, time, cost"],
        ["DishRating", "Stores meal feedback", "email, day, meal, dish, rating, date"],
        ["ReminderLog", "Prevents repeated scheduled processing", "key, status, recipients, sentAt"],
    ], [1.2, 2.0, 3.0])
    doc.add_heading("3.4 Main Modules", level=2)
    add_bullets(doc, [
        "Authentication module: handles Google OAuth login and session management.",
        "Student module: exposes user data, booked meal status, QR code data, order creation, order verification, and dish rating APIs.",
        "Admin module: supports menu and timing updates, meal count analysis, QR scanning, reminder trigger, and rollover trigger.",
        "Payment module: creates Razorpay orders and verifies payment signatures before saving purchased meal bookings.",
        "AI meal selection module: uses Gemini API to select meals from the weekly menu according to the student's diet preferences.",
        "Notification module: sends reminder emails through SMTP using configurable environment variables.",
        "Scheduler module: checks IST time and runs reminder/rollover tasks once per weekly cycle."
    ])
    doc.add_heading("3.5 AI-Based Meal Selection Workflow", level=2)
    add_numbered(doc, [
        "The student enters diet preferences such as avoiding fried items, preferring light food, or avoiding selected dishes.",
        "The frontend sends the preferences and weekly menu to the backend Gemini route.",
        "The backend prepares a prompt and asks Gemini to return only valid JSON with true or false values for each meal.",
        "The returned JSON is used to auto-select meals that match the student's preference.",
        "The student can review or change the selection before making payment."
    ])
    doc.add_heading("3.6 Weekly Meal Purchase Workflow", level=2)
    add_numbered(doc, [
        "The student logs in through Google OAuth.",
        "The student selects required meals for the upcoming week instead of paying blindly for an entire semester.",
        "The frontend requests Razorpay order creation from the backend.",
        "After payment, Razorpay returns order ID, payment ID, and signature.",
        "The backend verifies the signature using the Razorpay secret.",
        "The selected meals are saved in the buyer's next-week booking object and bought is marked true."
    ])
    doc.add_heading("3.7 Reminder Workflow", level=2)
    doc.add_paragraph(
        "The reminder scheduler checks the current time in Asia/Kolkata. At Sunday midnight, it identifies registered users whose email is not marked as bought for the upcoming week. "
        "It sends each of them an email warning that they have not purchased next-week meals and must buy before Sunday midnight to avail mess service for the next week. A scheduler log prevents duplicate reminder execution for the same weekly cycle."
    )
    doc.add_heading("3.8 Weekly Rollover Workflow", level=2)
    doc.add_paragraph(
        "At Monday midnight IST, the rollover scheduler moves each buyer's next-week meal booking object into the active this-week booking object. It then clears the next-week object and resets bought to false. "
        "This fixes the lifecycle issue where purchased meals remained stuck under next instead of becoming valid for the current week."
    )
    add_report_image(doc, "report_assets/meal_booking_lifecycle_workflow.png", "Figure 3.2: Weekly meal booking lifecycle workflow")

    # Chapter 4
    chapter_title(doc, 4, "Simulation and Results")
    doc.add_heading("4.1 Development Environment", level=2)
    add_table(doc, ["Component", "Tool/Technology"], [
        ["Programming Language", "JavaScript"],
        ["Frontend Framework", "React"],
        ["Backend Framework", "Express.js"],
        ["Runtime", "Node.js"],
        ["Database", "MongoDB"],
        ["Authentication", "Passport Google OAuth 2.0"],
        ["Payment Gateway", "Razorpay"],
        ["AI API", "Google Gemini API"],
        ["Email", "Nodemailer SMTP"],
        ["Deployment Configuration", "Environment variables in config/config.env"],
    ], [2.0, 4.2])
    doc.add_heading("4.2 Functional Testing", level=2)
    add_table(doc, ["Test Case", "Expected Result", "Status"], [
        ["Login using Google account", "User is authenticated and session is created", "[FILL: Pass/Fail]"],
        ["View weekly menu", "Menu data is displayed from backend", "[FILL: Pass/Fail]"],
        ["Use AI meal selection", "Meals matching diet preferences are auto-selected", "[FILL: Pass/Fail]"],
        ["Create Razorpay order", "Order ID and amount are returned", "[FILL: Pass/Fail]"],
        ["Verify successful payment", "Selected meals are saved under next week", "[FILL: Pass/Fail]"],
        ["Check booked meal using QR data", "Valid active meal booking is consumed once", "[FILL: Pass/Fail]"],
        ["Send reminder email", "Non-purchasing users receive email reminder", "[FILL: Pass/Fail]"],
        ["Run weekly rollover", "next-week bookings move to this-week bookings and bought resets", "[FILL: Pass/Fail]"],
        ["Admin meal count view", "Aggregate meal counts are displayed", "[FILL: Pass/Fail]"],
    ], [2.1, 3.2, 0.9])
    doc.add_heading("4.3 User Interface Results", level=2)
    doc.add_paragraph("The following placeholders should be replaced with final screenshots captured from the running application.")
    add_placeholder_box(doc, "Login screen / Google authentication", "Figure 4.1: Login screen placeholder")
    add_placeholder_box(doc, "Student dashboard or menu page", "Figure 4.2: Student menu page placeholder")
    add_placeholder_box(doc, "AI diet preference input and auto-selected meals", "Figure 4.3: AI meal selection placeholder")
    add_placeholder_box(doc, "Buy meals page with weekly meal selection", "Figure 4.4: Weekly meal purchase page placeholder")
    add_placeholder_box(doc, "QR code page", "Figure 4.5: QR code page placeholder")
    add_placeholder_box(doc, "Admin panel with menu/timing controls", "Figure 4.6: Admin panel placeholder")
    doc.add_heading("4.4 AI Meal Selection Result", level=2)
    doc.add_paragraph(
        "The AI meal selection feature helps the student choose meals without checking every item manually. For example, if the student writes that fried items should be avoided, the Gemini API checks the weekly menu and returns a meal selection where unsuitable meals are marked false. "
        "The student still has control and can review the selected meals before payment."
    )
    add_placeholder_box(doc, "Gemini response or UI showing diet-based meal selection", "Figure 4.7: AI meal selection result placeholder")
    doc.add_heading("4.5 Reminder Email Result", level=2)
    doc.add_paragraph(
        "The reminder feature can be tested by configuring SMTP credentials and calling the admin endpoint for sending reminder mails. "
        "Users who have not purchased next-week meals receive an email explaining that they have not bought their meals and must buy before Sunday midnight."
    )
    add_placeholder_box(doc, "Email inbox screenshot showing weekly meal purchase reminder mail", "Figure 4.8: Reminder email placeholder")
    doc.add_heading("4.6 Weekly Rollover Result", level=2)
    doc.add_paragraph(
        "The weekly rollover feature was added to solve the issue where meals purchased for the next week stayed marked as next week even after that week arrived. "
        "After rollover, the next-week bookings become current-week valid bookings, the next-week data is cleared, and the user can purchase meals for the following week."
    )
    add_table(doc, ["Before Rollover", "After Rollover"], [
        ["bought = true", "bought = false"],
        ["next contains purchased meals", "next is cleared"],
        ["this contains previous active week", "this contains purchased meals for active week"],
    ], [3.0, 3.0])
    add_placeholder_box(doc, "Database or API response before and after rollover", "Figure 4.9: Rollover verification placeholder")
    doc.add_heading("4.7 Benefits for Mess Management", level=2)
    doc.add_paragraph(
        "The portal is useful for the mess administration, not only for students. Since students select their meals in advance, the mess gets a better idea of how many people are expected for breakfast, lunch, and dinner. "
        "This makes it easier to decide the amount of raw material, cooking quantity, and serving preparation. Food wastage becomes lower because unnecessary extra food can be avoided. Shortage also becomes less likely because popular meals can be prepared according to expected demand."
    )
    doc.add_paragraph(
        "Advance headcount also helps control overcrowding. If the mess knows the expected crowd and the demand for each meal item, it can plan serving counters and food quantity more properly. This improves the dining experience and makes mess resource management smoother."
    )
    doc.add_heading("4.8 Security and Reliability Considerations", level=2)
    add_bullets(doc, [
        "Protected admin APIs verify that the authenticated email matches the configured admin email.",
        "Payment verification uses Razorpay signature validation before saving selected meal bookings.",
        "Session data is stored through Mongo-backed sessions.",
        "Mail credentials and payment secrets are read from environment variables instead of frontend code.",
        "Scheduler logs prevent duplicate weekly reminder or rollover execution."
    ])

    # Chapter 5
    chapter_title(doc, 5, "Conclusion and Future Work")
    doc.add_heading("5.1 Conclusion", level=2)
    doc.add_paragraph(
        "The IIITL Mess Portal presents a practical way to improve mess management in a college environment. It reduces the dependency on one-time semester mess payment by allowing students to select and buy meals for the upcoming week. "
        "This helps students avoid paying for meals they may not consume. It also helps the mess because the administration can see the expected meal count before preparation."
    )
    doc.add_paragraph(
        "The project combines Google login, menu display, AI-based meal selection using Gemini API, online payment, QR verification, admin management, meal feedback, email reminders, and weekly booking rollover. "
        "The main benefit is that both sides gain from the system. Students get more control over their mess spending, and the mess gets better planning data. This can reduce food wastage, help avoid shortage of food items, and make crowd management easier during meal time."
    )
    doc.add_paragraph(
        "A significant outcome of the project is the weekly rollover mechanism. Without rollover, meals purchased for the next week remain in the next-week state and do not become valid when the week changes. "
        "The implemented scheduler resolves this by promoting next-week bookings into current-week valid bookings every Monday at midnight IST."
    )
    doc.add_heading("5.2 Limitations", level=2)
    add_bullets(doc, [
        "The system currently depends on external services such as Google OAuth, Razorpay, MongoDB, and SMTP.",
        "The current admin authorization model is based on a single configured admin email.",
        "Exact production readiness depends on secure environment configuration, HTTPS deployment, and reliable server uptime.",
        "Manual screenshot insertion is still required for the final submitted report."
    ])
    doc.add_heading("5.3 Future Work", level=2)
    add_bullets(doc, [
        "Add role-based access control for multiple administrators and mess staff.",
        "Add refund/cancellation workflows for exceptional cases.",
        "Add analytics for food wastage prediction and demand trends.",
        "Improve AI meal selection by adding saved personal diet profiles for each student.",
        "Add push notifications or WhatsApp reminders in addition to email.",
        "Add a mobile-first progressive web app experience for faster QR access.",
        "Add detailed audit logs for QR scans, payment events, and admin changes.",
        "Add automated tests for payment verification, rollover, and reminder logic."
    ])

    # Appendix
    doc.add_page_break()
    add_centered_paragraph(doc, "Appendix Title Here", size=16, bold=True, spacing_after=18)
    doc.add_heading("Appendix A: Screenshot Checklist", level=2)
    add_bullets(doc, [
        "[IMAGE PLACEHOLDER: Homepage/Login]",
        "[IMAGE PLACEHOLDER: Weekly Menu]",
        "[IMAGE PLACEHOLDER: AI Diet Preference Selection]",
        "[IMAGE PLACEHOLDER: Buy Meals / Weekly Selection]",
        "[IMAGE PLACEHOLDER: Razorpay Payment]",
        "[IMAGE PLACEHOLDER: QR Code]",
        "[IMAGE PLACEHOLDER: Scan QR]",
        "[IMAGE PLACEHOLDER: Admin Panel]",
        "[IMAGE PLACEHOLDER: Total Meals]",
        "[IMAGE PLACEHOLDER: Reminder Email]",
        "[IMAGE PLACEHOLDER: Rollover API Test]"
    ])
    doc.add_heading("Appendix B: Important API Endpoints", level=2)
    add_table(doc, ["Endpoint", "Method", "Purpose"], [
        ["/api/auth/google", "GET", "Start Google login"],
        ["/api/data/menu", "GET", "Fetch weekly menu"],
        ["/api/data/time", "GET", "Fetch meal timings and costs"],
        ["/api/user/data", "GET", "Fetch logged-in buyer data"],
        ["/api/user/createOrder", "POST", "Create Razorpay order"],
        ["/api/user/checkOrder", "POST", "Verify payment and save selected meals"],
        ["/api/user/checkCoupon", "POST", "Validate and consume booked meal"],
        ["/api/user/gemini/selectMeals", "POST", "Auto-select meals using diet preferences and weekly menu"],
        ["/api/admin/setMenu", "POST", "Update menu"],
        ["/api/admin/meals", "POST", "Get total meal counts"],
        ["/api/admin/sendCouponReminder", "POST", "Manually trigger reminder email"],
        ["/api/admin/rolloverCoupons", "POST", "Manually trigger weekly rollover"],
    ], [2.2, 1.0, 3.0])
    doc.add_heading("Appendix C: Environment Variables", level=2)
    add_table(doc, ["Variable", "Purpose", "Placeholder"], [
        ["MONGO_URI", "MongoDB connection string", "[FILL]"],
        ["GOOGLE_CLIENT_ID", "Google OAuth client ID", "[FILL]"],
        ["GOOGLE_CLIENT_SECRET", "Google OAuth client secret", "[FILL]"],
        ["CALLBACK_URL", "OAuth callback URL", "[FILL]"],
        ["PAY_ID", "Razorpay key ID", "[FILL]"],
        ["PAY_SECRET", "Razorpay secret", "[FILL]"],
        ["GEMINI_API_KEY", "Gemini API key", "[FILL]"],
        ["ADMIN", "Admin email", "[FILL]"],
        ["MAIL_HOST", "SMTP host", "smtp.gmail.com"],
        ["MAIL_PORT", "SMTP port", "587"],
        ["MAIL_USER", "SMTP email username", "[FILL]"],
        ["MAIL_PASS", "SMTP app password", "[FILL]"],
        ["MAIL_FROM", "Sender name and address", "[FILL]"],
    ], [1.7, 2.7, 1.8])

    # Bibliography
    doc.add_page_break()
    add_centered_paragraph(doc, "Bibliography", size=16, bold=True, spacing_after=18)
    refs = [
        "React Documentation. React: The library for web and native user interfaces.",
        "Express.js Documentation. Fast, unopinionated, minimalist web framework for Node.js.",
        "MongoDB Documentation. MongoDB Manual and Mongoose ODM documentation.",
        "Passport.js Documentation. Authentication middleware for Node.js.",
        "Razorpay Documentation. Orders API and payment signature verification.",
        "Google Gemini API Documentation. Generative AI model usage for application features.",
        "Nodemailer Documentation. SMTP email sending for Node.js applications.",
        "OWASP Foundation. Web application security testing and secure development guidance.",
        "Google Identity Documentation. OAuth 2.0 authentication and authorization."
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.add_run("[{}] ".format(idx)).bold = True
        p.add_run(ref)

    # Footer on sections created by Word's continuous/page breaks remains linked; ensure present.
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        if not section.footer.paragraphs[0].text:
            add_footer_page_number(section)

    doc.save(OUT)


if __name__ == "__main__":
    build()
